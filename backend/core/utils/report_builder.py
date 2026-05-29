"""
core/utils/report_builder.py
============================
Constructs a .docx Word document from the resolved context + LLM-generated sections.

Document structure
------------------
  Title page
    ├── Report title
    ├── Resolution symbol(s) + date(s)
    └── Generated timestamp

  Section 1 — Resolution Overview & Analysis          [all / analyze]
    ├── Resolution detail table (symbol, title, date, body, vote counts)
    └── LLM narrative

  Section 2 — Voting Behavior                         [all / compare]
    ├── Country vote table (name, vote, blocs)
    └── LLM per-country analysis

  Section 3 — Bloc Alignments                         [all / blocs]
    ├── Bloc summary table
    └── LLM bloc analysis

  Section 4 — Voting Trends Over Time                 [all / timeline]
    ├── Multi-resolution vote matrix table
    └── LLM trend analysis

  Section 5 — Key Themes and Topics                   [all / themes]
    ├── Tag list
    └── LLM thematic analysis

  Section 6 — Custom Analysis                         [custom]
    └── LLM response to free-form question / country report

Returns an io.BytesIO containing the .docx bytes ready for HTTP response.
"""

from __future__ import annotations

import io
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from core.utils.report_context import BlocRow, CountryVoteRow, ReportContext, ResolutionMeta

# ── Colour palette ────────────────────────────────────────────────────────────
_DARK_BLUE  = RGBColor(0x1A, 0x23, 0x4E)   # title / headings
_ACCENT     = RGBColor(0xE6, 0x7E, 0x22)   # GeoStance orange
_YES_BG     = RGBColor(0xD4, 0xED, 0xDA)   # light green
_NO_BG      = RGBColor(0xF8, 0xD7, 0xDA)   # light red
_ABS_BG     = RGBColor(0xFF, 0xF3, 0xCD)   # light yellow
_HEADER_BG  = RGBColor(0x1A, 0x23, 0x4E)   # table header — dark blue
_HEADER_FG  = RGBColor(0xFF, 0xFF, 0xFF)   # white text on dark header

_VOTE_COLOURS = {
    "yes":        _YES_BG,
    "no":         _NO_BG,
    "abstain":    _ABS_BG,
    "absent":     RGBColor(0xE2, 0xE3, 0xE5),
    "not_member": RGBColor(0xE2, 0xE3, 0xE5),
    "—":          None,
}


# ── Low-level helpers ─────────────────────────────────────────────────────────

def _set_cell_bg(cell, colour: RGBColor | None) -> None:
    if colour is None:
        return
    from docx.oxml import OxmlElement
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), f"{colour[0]:02X}{colour[1]:02X}{colour[2]:02X}")
    tc_pr.append(shd)


def _bold_run(para, text: str, colour: RGBColor | None = None) -> None:
    run = para.add_run(text)
    run.bold = True
    if colour:
        run.font.color.rgb = colour


def _heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = _DARK_BLUE


def _table_header_row(table, headers: list[str]) -> None:
    row = table.rows[0]
    for i, header in enumerate(headers):
        cell = row.cells[i]
        cell.text = ""
        para = cell.paragraphs[0]
        run  = para.add_run(header)
        run.bold = True
        run.font.color.rgb = _HEADER_FG
        _set_cell_bg(cell, _HEADER_BG)


def _add_table(
    doc: Document,
    headers: list[str],
    rows:    list[list[str]],
    col_widths: list[float] | None = None,
    row_colours: list[RGBColor | None] | None = None,
) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    _table_header_row(table, headers)

    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg  = row_colours[r_idx] if row_colours else None
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(cell_text)
            if bg:
                _set_cell_bg(cell, bg)
            # Right-align numeric-looking cells
            if str(cell_text).replace("%", "").replace(".", "").isdigit():
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Set column widths if provided
    if col_widths:
        for col_idx, width in enumerate(col_widths):
            for cell in table.columns[col_idx].cells:
                cell.width = Inches(width)


def _vote_label(vote: str) -> str:
    return {
        "yes":        "In Favour",
        "no":         "Against",
        "abstain":    "Abstaining",
        "absent":     "Absent",
        "not_member": "Not a Member",
        "—":          "—",
    }.get(vote.lower(), vote.title())


def _vote_colour(vote: str) -> RGBColor | None:
    return _VOTE_COLOURS.get(vote.lower())


# ── Section builders ──────────────────────────────────────────────────────────

def _add_title_page(doc: Document, ctx: ReportContext, feature: str) -> None:
    doc.add_paragraph()

    # GeoStance brand line
    brand = doc.add_paragraph()
    brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = brand.add_run("GeoStance Intelligence Platform")
    run.font.size = Pt(11)
    run.font.color.rgb = _ACCENT
    run.bold = True

    doc.add_paragraph()

    # Report title
    title_map = {
        "all":      "UN Resolution Analysis Report",
        "analyze":  "Resolution Overview Report",
        "compare":  "Voting Behavior Report",
        "blocs":    "Bloc Alignment Report",
        "timeline": "Voting Trends Report",
        "themes":   "Key Themes Report",
        "custom":   "Custom Resolution Analysis",
    }
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title_map.get(feature, "Resolution Report"))
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = _DARK_BLUE
    title_run.bold = True

    doc.add_paragraph()

    # Resolution details
    for res in ctx.resolutions:
        sym_para = doc.add_paragraph()
        sym_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _bold_run(sym_para, res.un_symbol, _ACCENT)
        sym_para.add_run(f"  ·  {res.title[:70]}")
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para.add_run(f"{res.vote_date}  |  {res.body}")

    doc.add_paragraph()

    # Generated timestamp
    ts_para = doc.add_paragraph()
    ts_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ts_para.add_run(
        f"Generated: {datetime.utcnow().strftime('%d %B %Y, %H:%M UTC')}"
    ).font.color.rgb = RGBColor(0x6C, 0x75, 0x7D)

    doc.add_page_break()


def _add_overview_section(doc: Document, ctx: ReportContext, llm_text: str) -> None:
    _heading(doc, "1. Resolution Overview & Analysis", level=1)

    # Resolution detail table
    headers = ["Field", "Value"]
    rows = []
    for res in ctx.resolutions:
        rows += [
            ["Symbol",      res.un_symbol],
            ["Title",       res.title],
            ["Date",        res.vote_date],
            ["Body",        res.body],
            ["Event",       res.event_title],
            ["In Favour",   str(res.votes_yes)],
            ["Against",     str(res.votes_no)],
            ["Abstaining",  str(res.votes_abstain)],
            ["Absent",      str(res.votes_absent)],
        ]
        if len(ctx.resolutions) > 1:
            rows.append(["─────", "─────"])

    _add_table(doc, headers, rows, col_widths=[1.5, 5.0])
    doc.add_paragraph()
    doc.add_paragraph(llm_text)
    doc.add_page_break()


def _add_voting_behavior_section(doc: Document, ctx: ReportContext, llm_text: str) -> None:
    _heading(doc, "2. Voting Behavior", level=1)
    doc.add_paragraph(
        f"Analysis of the top {len(ctx.country_rows)} countries by geopolitical significance."
    ).italic = True

    # Country vote table
    if ctx.is_multi:
        symbols  = [r.un_symbol for r in ctx.resolutions]
        headers  = ["Country", "Blocs"] + symbols
        rows     = []
        colours  = []
        for cr in ctx.country_rows:
            row    = [cr.name, ", ".join(cr.blocs[:2]) or "—"]
            row   += [_vote_label(cr.multi_votes.get(sym, "—")) for sym in symbols]
            # Colour by first resolution's vote
            first_vote = cr.multi_votes.get(symbols[0], "—").lower()
            rows.append(row)
            colours.append(_vote_colour(first_vote))
    else:
        headers = ["Country", "ISO-3", "Vote", "Blocs"]
        rows    = []
        colours = []
        for cr in ctx.country_rows:
            rows.append([
                cr.name,
                cr.iso3,
                _vote_label(cr.vote),
                ", ".join(cr.blocs[:3]) or "—",
            ])
            colours.append(_vote_colour(cr.vote))

    _add_table(doc, headers, rows, row_colours=colours)
    doc.add_paragraph()
    doc.add_paragraph(llm_text)
    doc.add_page_break()


def _add_bloc_section(doc: Document, ctx: ReportContext, llm_text: str) -> None:
    _heading(doc, "3. Bloc Alignments", level=1)

    headers = ["Bloc", "Voted", "In Favour", "Against", "Abstaining", "% In Favour"]
    rows    = []
    colours = []
    for b in ctx.bloc_rows:
        voted = b.yes + b.no + b.abstain
        rows.append([
            b.name,
            str(voted),
            str(b.yes),
            str(b.no),
            str(b.abstain),
            f"{b.pct_yes}%",
        ])
        # Colour row by majority position
        if b.pct_yes >= 66:
            colours.append(_YES_BG)
        elif b.no > b.yes:
            colours.append(_NO_BG)
        elif b.abstain >= b.yes:
            colours.append(_ABS_BG)
        else:
            colours.append(None)

    _add_table(doc, headers, rows, row_colours=colours)
    doc.add_paragraph()
    doc.add_paragraph(llm_text)
    doc.add_page_break()


def _add_trends_section(doc: Document, ctx: ReportContext, llm_text: str) -> None:
    _heading(doc, "4. Voting Trends Over Time", level=1)

    if not ctx.is_multi:
        doc.add_paragraph("Select multiple resolutions to enable trend analysis.")
        doc.add_page_break()
        return

    symbols = [r.un_symbol for r in ctx.resolutions]
    headers = ["Country", "Blocs"] + symbols + ["Trend"]
    rows    = []
    colours = []

    _TREND = {"yes": 2, "abstain": 1, "no": 0, "absent": -1, "—": -1}

    for cr in ctx.country_rows:
        votes_seq = [cr.multi_votes.get(sym, "—").lower() for sym in symbols]
        scores    = [_TREND.get(v, -1) for v in votes_seq if v != "—"]
        if len(scores) >= 2:
            if scores[-1] > scores[0]:
                trend = "↑ Moving Toward Yes"
            elif scores[-1] < scores[0]:
                trend = "↓ Moving Toward No"
            else:
                trend = "→ Stable"
        else:
            trend = "—"

        row = [cr.name, ", ".join(cr.blocs[:2]) or "—"]
        row += [_vote_label(cr.multi_votes.get(sym, "—")) for sym in symbols]
        row += [trend]
        rows.append(row)
        first_vote = votes_seq[0] if votes_seq else "—"
        colours.append(_vote_colour(first_vote))

    _add_table(doc, headers, rows, row_colours=colours)
    doc.add_paragraph()
    doc.add_paragraph(llm_text)
    doc.add_page_break()


def _add_themes_section(doc: Document, ctx: ReportContext, llm_text: str) -> None:
    _heading(doc, "5. Key Themes and Topics", level=1)

    # Tag cloud as a simple table
    tags = ctx.all_tags[:20]
    if tags:
        tag_para = doc.add_paragraph()
        for i, tag in enumerate(tags):
            if i:
                tag_para.add_run("  ·  ")
            _bold_run(tag_para, tag.upper())

    doc.add_paragraph()
    doc.add_paragraph(llm_text)
    doc.add_page_break()


def _add_custom_section(doc: Document, ctx: ReportContext, llm_text: str, question: str) -> None:
    _heading(doc, "Custom Analysis", level=1)
    doc.add_paragraph(f'Query: "{question}"').italic = True
    doc.add_paragraph()
    doc.add_paragraph(llm_text)
    doc.add_page_break()


# ── Public API ────────────────────────────────────────────────────────────────

def build_docx(
    ctx:      ReportContext,
    sections: dict[str, str],
    feature:  str,
    question: str = "",
) -> io.BytesIO:
    """
    Build a .docx document from context + LLM-generated section texts.

    Args:
        ctx:      Assembled ReportContext (from report_context.py).
        sections: Dict of section_key → LLM text (from report_llm.py).
        feature:  "all"|"analyze"|"compare"|"blocs"|"timeline"|"themes"|"custom"
        question: Original free-form question (for "custom" feature title).

    Returns:
        BytesIO containing the finished .docx file.
    """
    doc = Document()

    # ── Page margins ─────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # ── Default body font ────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ── Title page ───────────────────────────────────────────────────────────
    _add_title_page(doc, ctx, feature)

    # ── Sections ─────────────────────────────────────────────────────────────
    if "overview" in sections:
        _add_overview_section(doc, ctx, sections["overview"])

    if "behavior" in sections:
        _add_voting_behavior_section(doc, ctx, sections["behavior"])

    if "blocs" in sections:
        _add_bloc_section(doc, ctx, sections["blocs"])

    if "trends" in sections:
        _add_trends_section(doc, ctx, sections["trends"])

    if "themes" in sections:
        _add_themes_section(doc, ctx, sections["themes"])

    if "custom" in sections:
        _add_custom_section(doc, ctx, sections["custom"], question)

    # ── Footer note ──────────────────────────────────────────────────────────
    footer_para = doc.add_paragraph(
        "Data sourced from the UN Digital Library (digitallibrary.un.org). "
        "AI analysis generated by GeoStance Intelligence Platform. "
        "Verify critical information before use."
    )
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer_para.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x6C, 0x75, 0x7D)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
